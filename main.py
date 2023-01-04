# 这是一个示例 Python 脚本。
# 运行pip install python-docx
import os

from PIL import Image
from docx import Document
from docx.shared import Inches


# 按 ⌃R 执行或将其替换为您的代码。
# 按 双击 ⇧ 在所有地方搜索类、文件、工具窗口、操作和设置。


def print_hi(name):
    # 在下面的代码行中使用断点来调试脚本。
    print(f'Hi, {name}')  # 按 ⌘F8 切换断点。


def picture_docx(path_picture, path_docx, name_docx):
    # 要插入的图片所在的文件夹
    # fold = 'C:\\Users\\Administrator\\Desktop\\tu'
    fold = path_picture
    if not os.path.exists(path_docx):
        os.makedirs(path_docx)
    # os.walk(fold)没有返回值，所以这么做显然没有结果，是错的
    # pics=list(os.walk(fold)[3])
    # # pics.pop()
    # print(pics)

    # pics是图片的名字
    # root是string类型， dirs和pics是list类型
    for root, dirs, pics in os.walk(fold):
        doc = Document()
        pics.sort()
        for i in range(0, len(pics)):
            print("开始插入图片" + pics[i])
            name = os.path.splitext(pics[i])[0]
            if name.startswith("."):
                continue
            # 我前半部分的路径直接复制黏贴了，没用root和dirs
            filepath = path_picture + pics[i]
            # filepath = root + '\\' + str(pics[i])
            doc.add_paragraph(name)

            try:
                doc.add_picture(filepath, width=Inches(6), height=Inches(8))
            except Exception as e:
                print("图片插入异常：" + e.__str__())
                pic_tmp = Image.open(filepath)
                # 如果格式有问题，就用save转换成默认的jpg格式
                pic_tmp.save(pic_tmp)
                # 把处理后的图片放进Document变量doc中
                doc.add_picture(pic_tmp, width=Inches(6), height=Inches(8))

            doc.add_page_break()
            # 把Document变量doc保存到指定路径的docx文件中
            doc.save(path_docx + name_docx + ".docx")

            # 输出保存成功的标志
            print("图片" + name, i + 1, "successfully added.", '\n')
    print("word文件处理完成。")
# 按间距中的绿色按钮以运行脚本。
if __name__ == '__main__':
    print_hi('PyCharm')
    picture_docx("/Users/mamenglong/Downloads/拉社保/", "doc/", "hhh")

# 访问 https://www.jetbrains.com/help/pycharm/ 获取 PyCharm 帮助
